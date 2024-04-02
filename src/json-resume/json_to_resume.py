import json
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

def generate_word_resume(data, filename):
    doc = Document()

    # Set margin sizes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    page_width = doc.sections[0].page_width
    left_margin = doc.sections[0].left_margin
    right_margin = doc.sections[0].right_margin
    tab_stop_position = page_width - right_margin*2

    # Personal Information
    name_heading = doc.add_heading(data['name'], 0)
    name_heading.paragraph_format.space_before = Pt(0)
    name_heading.paragraph_format.space_after = Pt(6)
    name_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Email: {data['email']} | Phone: {data['phone']}")
    doc.add_paragraph(f"Location: {data['location']} | LinkedIn: {data['linkedin']} | GitHub: {data['github']}")

    # Summary
    summary_heading = doc.add_heading('Summary', level=1)
    summary_heading.paragraph_format.space_before = Pt(6)
    summary_heading.paragraph_format.space_after = Pt(3)
    doc.add_paragraph(data['summary'])

    # Education
    edu_heading = doc.add_heading('Education', level=1)
    edu_heading.paragraph_format.space_before = Pt(6)
    edu_heading.paragraph_format.space_after = Pt(3)
    for edu in data['education']:
        doc.add_paragraph(f"{edu['degree']}, {edu['minor']}", style='List Bullet')
        doc.add_paragraph(f"{edu['institution']}, {edu['location']} | Graduation Year: {edu['duration']} | GPA: {edu['gpa']}")
        doc.add_paragraph(edu['details'])

    # Skills
    skills_heading = doc.add_heading('Skills', level=1)
    skills_heading.paragraph_format.space_before = Pt(6)
    skills_heading.paragraph_format.space_after = Pt(3)
    for skill_type, skills in data['skills'].items():
        skill_para = doc.add_paragraph(f"{skill_type.capitalize()}: {', '.join(skills)}")
        skill_para.paragraph_format.space_after = Pt(3)  # Adjust the spacing between skills

    # Experience
    exp_heading = doc.add_heading('Experience', level=1)
    exp_heading.paragraph_format.space_before = Pt(6)
    exp_heading.paragraph_format.space_after = Pt(3)

    # Set tab stops for experience section
    exp_tab_stops = doc.styles['Normal'].paragraph_format.tab_stops
    exp_tab_stops.add_tab_stop(tab_stop_position, alignment=WD_TAB_ALIGNMENT.RIGHT)

    for exp in data['experience']:


        
        company_duration_para = doc.add_paragraph(exp["company"])
        company_run = company_duration_para.add_run('\t')
        company_run.bold = True
        duration_run = company_duration_para.add_run(exp["duration"])
        duration_run.bold = True
        company_duration_para.paragraph_format.space_after = Pt(3)

        title_location_para = doc.add_paragraph(exp["title"])
        tile_run = title_location_para.add_run('\t')
        tile_run.italic = True
        location_run = title_location_para.add_run(exp["location"])
        location_run.italic = True
        title_location_para.paragraph_format.space_after = Pt(3)
        
        for detail in exp['details']:
            detail_para = doc.add_paragraph(detail, style='List Bullet')
            detail_para.paragraph_format.left_indent = Inches(0.25)
            detail_para.paragraph_format.space_after = Pt(3)

    # Projects
    projects_heading = doc.add_heading('Projects', level=1)
    projects_heading.paragraph_format.space_before = Pt(6)
    projects_heading.paragraph_format.space_after = Pt(3)

    # Set tab stops for projects section
    projects_tab_stops = doc.styles['Normal'].paragraph_format.tab_stops
    projects_tab_stops.add_tab_stop(tab_stop_position, alignment=WD_TAB_ALIGNMENT.RIGHT)

    for project in data['projects']:
        project_role = f"{project['title']} | {project['role']}"
        duration_url = f"{project['duration']} | {project['url']}"
        
        project_role_para = doc.add_paragraph(project_role)
        project_role_para.add_run('\t')
        project_role_para.add_run(duration_url)
        project_role_para.paragraph_format.space_after = Pt(3)
        
        for detail in project['details']:
            detail_para = doc.add_paragraph(detail, style='List Bullet')
            detail_para.paragraph_format.left_indent = Inches(0.25)
            detail_para.paragraph_format.space_after = Pt(3)

    # Certifications
    cert_heading = doc.add_heading('Certifications', level=1)
    cert_heading.paragraph_format.space_before = Pt(6)
    cert_heading.paragraph_format.space_after = Pt(3)
    for cert in data['certifications']:
        doc.add_paragraph(f"{cert['name']} | {cert['issuer']} | {cert['year']}", style='List Bullet')

    # Courses
    courses_heading = doc.add_heading('Courses', level=1)
    courses_heading.paragraph_format.space_before = Pt(6)
    courses_heading.paragraph_format.space_after = Pt(3)
    for course in data['courses']:
        doc.add_paragraph(f"{course['institution']} | {course['year']}", style='List Bullet')
        for name in course['name']:
            doc.add_paragraph(name, style='List Bullet 2')

    doc.save(filename)
if __name__ == '__main__':
    # Read JSON file
    fn = 'sample-resume.json'
    with open(fn) as file:
        data = json.load(file)

    import time 

    u = round(time.time())
    # Generate Word resume
    generate_word_resume(data, f'resume-{u}.docx')
